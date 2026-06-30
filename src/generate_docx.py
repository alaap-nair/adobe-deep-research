"""
generate_docx.py -- Export Assignment 7 deliverables to a compact Word document.

Usage:
    python src/generate_docx.py
"""

import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ANSWERS_DIR = os.path.join(ROOT, "outputs", "answers")


def slugify(q):
    slug = q.strip().lower()
    slug = re.sub(r"[^a-z0-9]+", "_", slug)
    return slug.strip("_")[:80].rstrip("_")


def set_cell(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.space_after = Pt(2)
    p.space_before = Pt(2)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True


def plain_table(doc, rows, cols):
    """Create a plain black-border table with no color fills."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    p.space_after = Pt(4)
    return p


def interpret_score(metric, score, question):
    """Generate a brief human-readable interpretation of a RAGAS score."""
    q = question.lower()

    if metric == "faithfulness":
        if score >= 0.9:
            return "Stuck to the context, did not hallucinate or make up facts"
        elif score >= 0.5:
            return "Mostly grounded but added some info not in the retrieved context"
        else:
            return "Did not stick to context — said 'I don't know' despite having relevant context retrieved"

    elif metric == "answer_relevancy":
        if score >= 0.8:
            return "Answer is responsive and directly addresses the question"
        elif score >= 0.5:
            return "Partially relevant — answered part of the question but missed key aspects"
        elif score > 0:
            return "Answer is only loosely related to what was asked"
        else:
            if "crispr" in q:
                return "Not relevant — returned 'I don't know' because CRISPR is not in our source data"
            else:
                return "Not relevant — returned 'I don't know' despite context being available"

    elif metric == "context_recall":
        if score >= 0.9:
            return "Retrieved the right chunks from the PDFs to answer the question"
        elif score >= 0.5:
            return "Found some relevant context but missed important chunks"
        else:
            if "crispr" in q:
                return "No relevant context exists — CRISPR not in source textbooks"
            else:
                return "Failed to retrieve the relevant chunks needed to answer"

    elif metric == "context_precision":
        if score >= 0.8:
            return "Retrieved context was clean and relevant, not much noise"
        elif score >= 0.5:
            return "Some noise in retrieved context but overall usable"
        elif score > 0:
            return "Lots of noise — retrieved irrelevant chunks alongside relevant ones"
        else:
            if "crispr" in q:
                return "No relevant context to retrieve"
            else:
                return "Retrieved context was mostly noise"

    return ""


def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── TITLE ──
    h = doc.add_heading("Assignment 7 — Deliverables", level=1)
    add_para(doc, "Team 2 | Molecular Biology & Genetics", size=11)

    # ── WORKSTREAM 1 ──
    doc.add_heading("Workstream 1: RAGAS Evaluation", level=2)

    # Load RAGAS results
    ragas_path = os.path.join(ROOT, "outputs", "ragas_results.json")
    with open(ragas_path) as f:
        ragas = json.load(f)

    pqs = ragas.get("per_question", [])
    avg = ragas.get("averages", {})

    # Numeric scores table
    add_para(doc, "Scores", bold=True, size=11)
    table = plain_table(doc, rows=1, cols=5)
    for i, h_text in enumerate(["Question", "Faithfulness", "Answer Relevancy", "Context Recall", "Context Precision"]):
        set_cell(table.rows[0].cells[i], h_text, bold=True, size=9)

    for pq in pqs:
        row = table.add_row()
        q = pq["question"]
        if len(q) > 45:
            q = q[:43] + "..."
        set_cell(row.cells[0], q, size=9)
        set_cell(row.cells[1], f"{pq['faithfulness']:.2f}", size=9)
        set_cell(row.cells[2], f"{pq['answer_relevancy']:.2f}", size=9)
        set_cell(row.cells[3], f"{pq['context_recall']:.2f}", size=9)
        set_cell(row.cells[4], f"{pq['context_precision']:.2f}", size=9)

    row = table.add_row()
    set_cell(row.cells[0], "AVERAGE", bold=True, size=9)
    set_cell(row.cells[1], f"{avg.get('faithfulness', 0):.2f}", bold=True, size=9)
    set_cell(row.cells[2], f"{avg.get('answer_relevancy', 0):.2f}", bold=True, size=9)
    set_cell(row.cells[3], f"{avg.get('context_recall', 0):.2f}", bold=True, size=9)
    set_cell(row.cells[4], f"{avg.get('context_precision', 0):.2f}", bold=True, size=9)

    doc.add_paragraph("")

    # Human-readable interpretation table (matches the screenshot style)
    add_para(doc, "Interpretations", bold=True, size=11)
    table = plain_table(doc, rows=1, cols=4)
    for i, h_text in enumerate(["Faithfulness", "Answer Relevancy", "Context Recall", "Context Precision"]):
        set_cell(table.rows[0].cells[i], h_text, bold=True, size=9)

    for pq in pqs:
        row = table.add_row()
        for i, metric in enumerate(["faithfulness", "answer_relevancy", "context_recall", "context_precision"]):
            interp = interpret_score(metric, pq[metric], pq["question"])
            set_cell(row.cells[i], interp, size=9)

    doc.add_paragraph("")

    # Brief summary
    add_para(doc,
        "Faithfulness is strong (0.89) — the system does not hallucinate. "
        "Answer relevancy (0.49) and context recall (0.56) are the weak points: "
        "CRISPR is absent from our source data entirely (0.0 across all metrics for both CRISPR questions), "
        "and broad queries like 'What is genetics?' fail at entity resolution, causing the system to say "
        "'I don't know' despite having relevant context. Multi-hop questions suffer from noisy retrieval "
        "(context precision 0.21 for RNA regulation).",
        size=10,
    )

    # ── WORKSTREAM 2 ──
    doc.add_heading("Workstream 2: Expanded Dataset (20 Questions)", level=2)

    with open(os.path.join(ROOT, "data", "ground_truth.json")) as f:
        questions = json.load(f)

    # Original 10 system outputs + ratings from user
    original_data = {
        "What is the primary function of DNA?": ("DNA encodes RNA; RNA encodes protein.", 4, "Minor — could mention central dogma more explicitly"),
        "What is the primary function of RNA?": ("RNA encodes protein.", 4, "Partial — omits intermediary function between DNA and protein"),
        "What is CRISPR?": ("I don't know based on the provided context.", 5, "Correct refusal — CRISPR not in source textbooks"),
        "How does protein synthesis work?": ("Protein synthesis involves the decoding of an mRNA message into a polypeptide product.", 5, "Accurate and comprehensive"),
        "What is a genome?": ("The complete set of genes.", 4, "Minor omission — missing 'genetic instructions found in a cell'"),
        "How does RNA regulate gene expression?": ("Small nuclear RNAs have a variety of functions, including splicing pre-mRNAs and regulating transcription factors.", 3, "Only covers snRNA role, misses inhibition/degradation"),
        "What are the differences between the codons in DNA and RNA?": ("The only nucleotide difference is that in mRNA, all of the T nucleotides are replaced with U nucleotides.", 5, "Accurate — directly addresses thymine vs uracil"),
        "Why is RNA only a single helix?": ("ERROR — LLM returned malformed JSON, pipeline crashed.", 1, "Pipeline crash — missing 'reasoning' field in LLM response"),
        "What is genetics?": ("I don't know based on the provided context.", 1, "False negative — context was present but entity resolution failed on broad term"),
        "How does CRISPR work?": ("I don't know based on the provided context.", 1, "Correct refusal — CRISPR/Cas9 not in any source document"),
    }

    # Q&A table
    table = plain_table(doc, rows=1, cols=6)
    for i, h_text in enumerate(["#", "Question", "Difficulty", "Ground Truth", "System Output", "Rating"]):
        set_cell(table.rows[0].cells[i], h_text, bold=True, size=8)

    for i, q in enumerate(questions, 1):
        row = table.add_row()
        question = q["question"]
        difficulty = q.get("difficulty", "Standard")
        gt = q["ground_truth"]

        # Get system output
        if question in original_data:
            output, rating, _ = original_data[question]
        else:
            slug = slugify(question)
            answer_path = os.path.join(ANSWERS_DIR, f"{slug}.json")
            if os.path.exists(answer_path):
                with open(answer_path) as f:
                    data = json.load(f)
                output = data.get("answer", "")
                rating = "TBD"
            else:
                output = "[Not yet run]"
                rating = "—"

        # No truncation — show full text

        set_cell(row.cells[0], str(i), size=8)
        set_cell(row.cells[1], question, size=8)
        set_cell(row.cells[2], difficulty, size=8)
        set_cell(row.cells[3], gt, size=8)
        set_cell(row.cells[4], output, size=8)
        set_cell(row.cells[5], str(rating), size=8)

    # Set column widths
    col_widths = [Inches(0.3), Inches(1.8), Inches(0.7), Inches(2.0), Inches(2.0), Inches(0.4)]
    for row_obj in table.rows:
        for i, cell in enumerate(row_obj.cells):
            cell.width = col_widths[i]

    doc.add_paragraph("")

    # Failure analysis for new questions (brief)
    add_para(doc, "Failure Analysis (New Questions)", bold=True, size=11)
    for q in questions:
        if q.get("set") != "new":
            continue
        targets = q.get("targets_failure", "")
        if targets:
            add_para(doc, f"Q: {q['question']}", bold=True, size=9)
            add_para(doc, f"Targets: {targets}", size=9)

    # ── WORKSTREAM 3A ──
    doc.add_heading("Workstream 3A: Deduplication & Canonicalization", level=2)

    add_para(doc, "Strategies Researched", bold=True)
    add_para(doc,
        "1. Embedding-based clustering — use BioBERT or BGE embeddings to cluster "
        "near-duplicate entities by cosine similarity. Pros: catches semantic duplicates. "
        "Cons: risk of over-merging, requires threshold tuning.",
        size=10,
    )
    add_para(doc,
        "2. Rule-based normalization (implemented) — deterministic pipeline: strip "
        "parenthetical abbreviations, strip leading articles, strip punctuation, "
        "filter noise entities. Pros: fast, no false merges. Cons: won't catch "
        "semantic duplicates.",
        size=10,
    )

    add_para(doc, "Before vs. After", bold=True)
    table = plain_table(doc, rows=1, cols=3)
    for i, h_text in enumerate(["Metric", "Before", "After"]):
        set_cell(table.rows[0].cells[i], h_text, bold=True, size=9)
    for metric, before, after in [
        ("Triples", "1,761", "1,745 (-16 noise removed)"),
        ("Unique entity IDs", "1,979", "1,957 (-22 merged/removed)"),
        ("Noise entities (single chars, fragments)", "22+", "0 (filtered)"),
        ("Parenthetical duplicates", "25+", "Merged (e.g. 'messenger RNA (mRNA)' = 'messenger RNA')"),
    ]:
        row = table.add_row()
        set_cell(row.cells[0], metric, size=9)
        set_cell(row.cells[1], before, size=9)
        set_cell(row.cells[2], after, size=9)

    doc.add_paragraph("")
    add_para(doc,
        "RAGAS impact: changes are prospective — future extractions produce cleaner graphs. "
        "Expected to improve context precision (currently 0.56) by reducing irrelevant entity matches.",
        size=10,
    )

    # ── WORKSTREAM 3B ──
    doc.add_heading("Workstream 3B: Schema Refinement", level=2)

    add_para(doc, "Schema Definition", bold=True)
    add_para(doc, "Node types: Gene, Protein, Molecule, CellStructure, BiologicalProcess, "
        "Pathway, Organism, Disease, Technique, GenomicElement", size=10)
    add_para(doc, "Relationship types: encodes, transcribes, translates, regulates, inhibits, "
        "activates, catalyzes, binds to, produces, degrades, contains, interacts with, "
        "causes, prevents, mutates, repairs, replicates, modifies, transports, "
        "is located in, is type of, requires, uses, converts (30 total)", size=10)

    add_para(doc, "Implementation", bold=True)
    add_para(doc,
        "Schema defined in src/bio_schema.py and appended to the LLM extraction prompt. "
        "The prompt now asks the LLM to classify each entity with a head_type/tail_type and "
        "instructs: 'Do NOT extract raw numbers, figure references, or generic phrases.'",
        size=10,
    )

    add_para(doc, "Coverage", bold=True)

    coverage_path = os.path.join(ROOT, "outputs", "coverage_report.json")
    with open(coverage_path) as f:
        coverage = json.load(f)

    add_para(doc,
        f"Overall: {coverage['covered_terms']}/{coverage['total_terms']} terms "
        f"({coverage['overall_coverage_pct']}%). "
        "Only gap: CRISPR and Cas9 (not in source textbooks).",
        size=10,
    )

    # Save
    output_path = os.path.join(ROOT, "outputs", "assignment7_deliverables.docx")
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == "__main__":
    main()
